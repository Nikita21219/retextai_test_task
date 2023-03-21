from docx import Document
import docx
from googletrans import Translator


class FileTranslator:
    def __init__(self):
        self._translator = Translator()
        self._dest = 'en'

    def _translate_text(self, text):
        return self._translator.translate(text, dest=self._dest).text

    def _translate_table(self, table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = self._translate_text(paragraph.text)

    def translate(self, input_path, output_path) -> bool:
        if not input_path or not output_path:
            return False
        try:
            doc = Document(input_path)
        except (ValueError, docx.opc.exceptions.PackageNotFoundError):
            return False
        for paragraph in doc.paragraphs:
            paragraph.text = self._translate_text(paragraph.text)

        for table in doc.tables:
            self._translate_table(table)

        doc.save(output_path)
        return True
